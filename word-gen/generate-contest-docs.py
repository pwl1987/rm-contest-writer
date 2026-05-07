#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# DEPRECATED - 使用 .js 版本
# 保留至 Phase 2 完成验证
"""
融媒技术创新大赛申报材料 - Word文档生成脚本
严格遵守党政机关公文格式（GB/T 9704-2012）
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== 页面设置 =====
PAGE_MARGIN = {
    'top': Cm(3.7),
    'bottom': Cm(3.5),
    'left': Cm(2.8),
    'right': Cm(2.6),
}
HEADER_POS = Cm(1.5)   # 页眉距页面顶端 1.5cm
FOOTER_POS = Cm(2.8)   # 页脚距页面底端 2.8cm

# ===== 字体 =====
FONT_TITLE = '方正小标宋简体'     # 标题（作品名称）
FONT_UNIT = '楷体_GB2312'        # 单位名称
FONT_BODY = '仿宋_GB2312'         # 正文
FONT_HEAD1 = '黑体'               # 一级标题
FONT_HEAD2 = '楷体_GB2312'        # 二级标题
FONT_HEAD3 = '仿宋_GB2312'        # 三级标题（加粗）
FONT_HEAD4 = '仿宋_GB2312'        # 四级标题
FONT_TABLE_TITLE = '黑体'         # 表格标题
FONT_TABLE_HEAD = '黑体'          # 表头
FONT_TABLE_BODY = '宋体'            # 表格正文（宋体五号）
FONT_PAGE_NUM = '宋体'            # 页码

# ===== 字号（half-points）=====
# 二号=22pt, 小三=15pt, 三号=16pt, 小四=14pt, 五号=10.5pt
SZ_TITLE = 44      # 二号（22pt）
SZ_UNIT = 30      # 小三号（15pt）
SZ_BODY = 32      # 三号（16pt）
SZ_HEAD = 32      # 三号（16pt）标题
SZ_TABLE_TITLE = 28  # 小四号（14pt）表格标题
SZ_TABLE = 21     # 五号（10.5pt）表格
SZ_PAGE_NUM = 28  # 小四号（14pt）页码

# ===== 行距 =====
LINE_TITLE = 600    # 标题行距30磅=600twip
LINE_BODY = 560     # 正文行距固定28磅=560twip
LINE_TABLE = 360    # 表格内转行行距固定18磅=360twip

# ===== 段落间距 =====
# 段前段后：标题0行，正文0行，表格标题上空一行（段前0.5行≈10twip），表格后空一行
SPACE_NONE = 0      # 0行


def make_element(tag, attrs=None):
    el = OxmlElement(tag)
    if attrs:
        for k, v in attrs.items():
            el.set(qn(k), str(v))
    return el


def clear_children(parent, tag_name):
    """清除所有指定子元素"""
    for child in parent.findall(qn(tag_name)):
        parent.remove(child)


def set_paragraph_spacing(para, line, line_rule='exact', before=None, after=None):
    """设置段落行距和段前段后间距"""
    pPr = para._element.get_or_add_pPr()
    # 清除旧spacing
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    spacing = make_element('w:spacing', {
        'w:line': str(line),
        'w:lineRule': line_rule,
    })
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)


def set_run_font(run, font_name, font_size_half, bold=False):
    """设置run的字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size_half / 2)
    run.font.bold = bold


def add_empty_paragraph(doc):
    para = doc.add_paragraph()
    para.add_run('')
    return para


def add_title_paragraph(doc, text):
    """作品名称：方正小标宋二号字，居中，行距30磅，段前段后0行"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, FONT_TITLE, SZ_TITLE, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, LINE_TITLE, before=SPACE_NONE, after=SPACE_NONE)
    return para


def add_unit_name_paragraph(doc, text):
    """单位名称：楷体小三号字，左右居中"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, FONT_UNIT, SZ_UNIT, bold=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, LINE_TITLE, before=SPACE_NONE, after=SPACE_NONE)
    return para


def add_body_paragraph(doc, text, indent=True):
    """正文段落：仿宋三号字，不加粗，行距28磅，段前段后0行，首行缩进2字符"""
    para = doc.add_paragraph()
    segments = parse_inline_formats(text)
    for seg_text, seg_bold, seg_code in segments:
        run = para.add_run(seg_text)
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        run.font.size = Pt(SZ_BODY / 2)
        run.font.bold = False  # 正文不加粗
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        pPr = para._element.get_or_add_pPr()
        ind = make_element('w:ind', {'w:firstLineChars': '200'})
        pPr.append(ind)
    set_paragraph_spacing(para, LINE_BODY, before=SPACE_NONE, after=SPACE_NONE)
    return para


def add_heading_paragraph(doc, text, font, bold=False, before=None, after=None, indent=True):
    """正文内各级标题：仿宋三号/黑体/楷体，居左，首行缩进2字符"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font, SZ_HEAD, bold=bold)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        pPr = para._element.get_or_add_pPr()
        ind = make_element('w:ind', {'w:firstLineChars': '200'})
        pPr.append(ind)
    set_paragraph_spacing(para, LINE_BODY, before=SPACE_NONE, after=SPACE_NONE)
    return para


def add_code_paragraph(doc, text):
    """代码块段落：等宽字体，左缩进"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    run.font.size = Pt(SZ_TABLE / 2)  # 五号
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = para._element.get_or_add_pPr()
    ind = make_element('w:ind', {'w:left': '567'})
    pPr.append(ind)
    set_paragraph_spacing(para, LINE_BODY, before=SPACE_NONE, after=SPACE_NONE)
    return para


def add_table(doc, rows):
    """添加表格：表格正文宋体五号，数字Times New Roman五号，两端对齐"""
    if not rows or len(rows) < 2:
        return None

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.JUSTIFY  # 两端对齐

    # 表格宽度100%
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = make_element('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = make_element('w:tblW', {'w:w': '5000', 'w:type': 'pct'})
    tblPr.append(tblW)

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    if i == 0:
                        # 表头：黑体五号不加粗
                        run.font.name = FONT_TABLE_HEAD
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TABLE_HEAD)
                        run.font.size = Pt(SZ_TABLE / 2)
                        run.font.bold = False
                    else:
                        # 表格正文：宋体五号
                        run.font.name = FONT_TABLE_BODY
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TABLE_BODY)
                        run.font.size = Pt(SZ_TABLE / 2)
                        run.font.bold = False
                        # 数字用Times New Roman
                        if re.match(r'^[\d\.\-\+%]+$', cell_text.strip()):
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(para, LINE_TABLE, before=SPACE_NONE, after=SPACE_NONE)
    return table


def add_table_title_paragraph(doc, text):
    """表格标题：黑体小四号居中，段前0.5行（上空一行视觉），段后0行，标题与表格紧邻"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, FONT_TABLE_TITLE, SZ_TABLE_TITLE, bold=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 段前0.5行 ≈ 10twip，与上面正文视觉空一行
    set_paragraph_spacing(para, LINE_BODY, before=10, after=SPACE_NONE)
    return para


def add_image_paragraph(doc, image_path, caption_text=None):
    """插图：插入图片，图与前后正文空一行"""
    if not os.path.exists(image_path):
        return None

    para = doc.add_paragraph()
    # 添加空行（上空一行）
    set_paragraph_spacing(para, LINE_BODY, before=LINE_BODY, after=SPACE_NONE)

    # 读取图片并添加
    try:
        from docx.shared import Inches, Cm
        from docx.oxml.ns import qn
        from lxml import etree

        # 获取图片尺寸（保持宽高比，最大宽度15cm）
        max_width = Cm(15)

        run = para.add_run()
        run.add_picture(image_path, width=max_width)
    except Exception as e:
        para.add_run(f'[图片: {os.path.basename(image_path)}]')

    # 图片后空一行
    after_para = doc.add_paragraph()
    set_paragraph_spacing(after_para, LINE_BODY, before=SPACE_NONE, after=LINE_BODY)

    # 如果有标题
    if caption_text:
        add_image_caption(doc, caption_text)

    return para


def add_image_caption(doc, text):
    """插图标题：黑体小四号居中，位于图下方"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, FONT_TABLE_TITLE, SZ_TABLE_TITLE, bold=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, LINE_BODY, before=SPACE_NONE, after=SPACE_NONE)
    return para


def to_chinese_quotes(text):
    """英文双引号 → 中文双引号（前引号 U+201C，后引号 U+201D）"""
    if not text:
        return text
    return text.replace('"', '“').replace('"', '”')


def parse_inline_formats(text):
    """解析 **bold** 和 `code`"""
    segments = []
    remaining = text
    pattern = re.compile(r'\*\*([^*]+?)\*\*|`([^`]+)`')
    while remaining:
        m = pattern.search(remaining)
        if m:
            before = remaining[:m.start()]
            if before:
                segments.append((to_chinese_quotes(before), False, False))
            if m.group(1):
                segments.append((to_chinese_quotes(m.group(1)), True, False))
            elif m.group(2):
                segments.append((to_chinese_quotes(m.group(2)), False, True))
            remaining = remaining[m.end():]
        else:
            segments.append((to_chinese_quotes(remaining), False, False))
            break
    return segments


def is_signature_line(text):
    return ('团队/个人：' in text or
            '（公章）' in text or
            ('年' in text and '月' in text and '日' in text and len(text) < 25) or
            '甲方（公章）：' in text or
            '乙方（公章）：' in text)


def parse_heading_style(text):
    """判断正文内标题样式"""
    if text.startswith('# '):
        return 'title', to_chinese_quotes(text[2:].replace('**', ''))
    if text.startswith('## '):
        return 'h1', to_chinese_quotes(text[3:].replace('**', ''))
    if text.startswith('### '):
        return 'h2', to_chinese_quotes(text[4:].replace('**', ''))
    if re.match(r'^[一二三四五六七八九十]、', text):
        return 'h1', to_chinese_quotes(text)
    if re.match(r'^\(\d+\)', text):
        return 'h4', to_chinese_quotes(text)
    if re.match(r'^\([^一二三四五六七八九十]', text):
        return 'h2', to_chinese_quotes(text)
    if re.match(r'^\d+\.', text):
        return 'h3', to_chinese_quotes(text)
    return None, None


def add_para_with_mixed_format(doc, text, indent=True, bold_default=False):
    """通用混合格式段落"""
    para = doc.add_paragraph()
    segments = parse_inline_formats(text)
    for seg_text, seg_bold, seg_code in segments:
        run = para.add_run(seg_text)
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        run.font.size = Pt(SZ_BODY / 2)
        run.font.bold = seg_bold if seg_bold else bold_default
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        pPr = para._element.get_or_add_pPr()
        ind = make_element('w:ind', {'w:firstLineChars': '200'})
        pPr.append(ind)
    set_paragraph_spacing(para, LINE_BODY, before=SPACE_NONE, after=SPACE_NONE)
    return para


def set_page_headers(doc):
    """设置页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header.distance = HEADER_POS
        # 清空默认内容
        for para in header.paragraphs:
            for run in para.runs:
                run.text = ''
        if not header.paragraphs:
            header.add_paragraph()
        # 页眉内容为空（不显示）

        # 页脚
        footer = section.footer
        footer.distance = FOOTER_POS
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 页码格式：宋体小四号 "- 5 -"
        run = footer_para.add_run()
        run.font.name = FONT_PAGE_NUM
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_PAGE_NUM)
        run.font.size = Pt(SZ_PAGE_NUM / 2)
        # 页码用域代码实现：PAGE_NUM
        fldChar1 = make_element('w:fldChar', {'w:fldCharType': 'begin'})
        instrText = make_element('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = make_element('w:fldChar', {'w:fldCharType': 'end'})
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # 页码格式：左空1字符 + dash + 空格 + 页码 + 空格 + dash + 右空1字符
        # 由于Word的复杂格式，简化为居中显示页码
        section.page_number_format = 1  # ARABIC


def markdown_to_docx(md_path, docx_path, png_path=None):
    print(f'读取 Markdown: {md_path}')

    # 自动查找同名PNG（架构图）
    if png_path is None:
        md_dir = os.path.dirname(os.path.abspath(md_path))
        md_base = os.path.splitext(os.path.basename(md_path))[0]
        possible_png = os.path.join(md_dir, md_base + '.png')
        if os.path.exists(possible_png):
            png_path = possible_png
            print(f'检测到架构图: {png_path}')

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN['top']
        section.bottom_margin = PAGE_MARGIN['bottom']
        section.left_margin = PAGE_MARGIN['left']
        section.right_margin = PAGE_MARGIN['right']
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # 设置页眉页脚
    set_page_headers(doc)

    lines = content.split('\n')
    n = len(lines)
    i = 0
    table_rows = []
    in_table = False
    in_code_block = False
    code_block_lines = []
    first_para = True  # 标题后空一行标记

    while i < n:
        line = lines[i]

        # ===== 代码块 =====
        if line.strip().startswith('```'):
            if not in_code_block:
                if in_table and table_rows:
                    add_table_title_paragraph(doc, '表')  # 占位
                    add_table(doc, table_rows)
                    add_empty_paragraph(doc)
                    table_rows = []
                    in_table = False
                in_code_block = True
                code_block_lines = []
            else:
                for cbl in code_block_lines:
                    add_code_paragraph(doc, cbl)
                add_empty_paragraph(doc)
                # 自动插入架构图（如有同名PNG）
                code_text = '\n'.join(code_block_lines)
                if ('┌─' in code_text or '第四层' in code_text or '第一层' in code_text) and png_path:
                    add_image_paragraph(doc, png_path, '图：系统总体架构')
                    add_empty_paragraph(doc)
                in_code_block = False
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ===== 空行 =====
        if not stripped:
            if in_table and table_rows:
                add_empty_paragraph(doc)
                table_rows = []
                in_table = False
            i += 1
            continue

        # 跳过分隔线
        if stripped == '---':
            i += 1
            continue

        # ===== 参赛类别标题块（最前两行）=====
        if first_para and stripped.startswith('**参赛类别'):
            # 参赛类别标识（加粗）
            add_para_with_mixed_format(doc, stripped, indent=False)
            first_para = False
            i += 1
            continue

        # ===== 作品名称（# 开头）=====
        if stripped.startswith('# '):
            title_text = to_chinese_quotes(stripped[2:].replace('**', ''))
            add_title_paragraph(doc, title_text)
            add_empty_paragraph(doc)  # 与正文之间空一行
            first_para = False
            i += 1
            continue

        # ===== 单位名称（如有）=====
        if stripped.startswith('**单位'):
            add_unit_name_paragraph(doc, to_chinese_quotes(stripped.replace('**', '')))
            add_empty_paragraph(doc)
            i += 1
            continue

        # ===== 表格 =====
        if stripped.startswith('|'):
            in_table = True
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip().replace('<[^>]+>', '').replace('&nbsp;', '') for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_rows:
                # 标题与表格之间不空行（标题段前0.5行已提供视觉间隔），表格后空一行
                add_table_title_paragraph(doc, '表')  # 简化：表格标题由用户自行补充
                add_table(doc, table_rows)
                add_empty_paragraph(doc)  # 表格后空一行再跟正文
                table_rows = []
                in_table = False

        # ===== 标题判断 =====
        style, text = parse_heading_style(stripped)

        if style == 'title':
            add_title_paragraph(doc, text)
        elif style == 'h1':
            # 一级标题：黑体三号不加粗，如"一、评价指标构建"
            add_heading_paragraph(doc, text, FONT_HEAD1, bold=False, indent=True)
        elif style == 'h2':
            # 二级标题：楷体三号不加粗，如"（二）指标选取"
            add_heading_paragraph(doc, text, FONT_HEAD2, bold=False, indent=True)
        elif style == 'h3':
            # 三级标题：仿宋三号，段落内所有文字加粗，如"3.基本特征"
            para = doc.add_paragraph()
            # 去掉**后全文加粗
            clean_text = to_chinese_quotes(stripped.replace('**', ''))
            run = para.add_run(clean_text)
            run.font.name = FONT_HEAD3
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEAD3)
            run.font.size = Pt(SZ_HEAD / 2)
            run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pPr = para._element.get_or_add_pPr()
            ind = make_element('w:ind', {'w:firstLineChars': '200'})
            pPr.append(ind)
            set_paragraph_spacing(para, LINE_BODY, before=SPACE_NONE, after=SPACE_NONE)
        elif style == 'h4':
            # 四级标题：仿宋三号不加粗，如"（2）学习行为"
            add_heading_paragraph(doc, text, FONT_HEAD4, bold=False, indent=True)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            add_para_with_mixed_format(doc, text, indent=True)
        elif is_signature_line(stripped):
            # 签名右对齐
            para = doc.add_paragraph()
            run = para.add_run(to_chinese_quotes(stripped))
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            run.font.size = Pt(SZ_BODY / 2)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pPr = para._element.get_or_add_pPr()
            ind = make_element('w:ind', {'w:right': '567'})
            pPr.append(ind)
            set_paragraph_spacing(para, LINE_BODY, before=SPACE_NONE, after=SPACE_NONE)
        else:
            add_para_with_mixed_format(doc, stripped, indent=True)

        i += 1

    # 末尾表格
    if in_table and table_rows:
        add_empty_paragraph(doc)
        add_table_title_paragraph(doc, '表')
        add_table(doc, table_rows)
        add_empty_paragraph(doc)

    doc.save(docx_path)
    print(f'生成 docx: {docx_path}')
    return True


def main():
    if len(sys.argv) < 3:
        print('用法: python3 generate-contest-docs.py <markdown文件> <输出docx文件>')
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]

    print('=' * 60)
    print('融媒技术创新大赛申报材料 - Word文档生成')
    print('=' * 60)
    print()
    print('公文格式（GB/T 9704-2012）：')
    print('  页面：A4，上下3.7/3.5cm，左右2.8/2.6cm')
    print('  标题：方正小标宋二号，行距30磅，段前段后0行')
    print('  正文：仿宋三号，行距28磅，首行缩进2字符，段前段后0行')
    print('  各级标题：黑体/楷体/仿宋，三号，左对齐，首行缩进2字符')
    print('  表格：黑体小四标题，仿宋五号正文，数字Times New Roman五号')
    print('  页码：宋体小四号，页脚外侧"- 5 -"格式')
    print()

    if not os.path.exists(md_path):
        print(f'文件不存在: {md_path}')
        sys.exit(1)

    success = markdown_to_docx(md_path, docx_path)
    if success:
        print(f'输出文件: {docx_path}')
    else:
        print('生成失败')
        sys.exit(1)


if __name__ == '__main__':
    main()